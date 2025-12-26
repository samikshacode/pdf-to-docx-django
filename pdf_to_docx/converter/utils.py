import pdfplumber
from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH

def pdf_to_docx(pdf_path, output_path):
    document = Document()

    def add_center(text, size=12, bold=False):
        p = document.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(text)
        run.bold = bold
        run.font.size = Pt(size)

    def add_left(text, size=11, bold=False):
        p = document.add_paragraph()
        run = p.add_run(text)
        run.bold = bold
        run.font.size = Pt(size)

    with pdfplumber.open(pdf_path) as pdf:
        for page in pdf.pages:
            text = page.extract_text()
            if not text:
                continue

            lines = text.split("\n")
            for line in lines:
                line = line.strip()

                if not line:
                    document.add_paragraph("")
                elif line.isupper():
                    add_center(line, bold=True)
                elif "DETAILS" in line or "FORM" in line:
                    add_left(line, bold=True)
                else:
                    add_left(line)

    document.save(output_path)
